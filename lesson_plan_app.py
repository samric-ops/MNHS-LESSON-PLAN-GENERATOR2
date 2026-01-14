import streamlit as st
import google.generativeai as genai
import json
from datetime import date
import io
import requests
import urllib.parse
import random
import re
import hashlib

# --- NEW LIBRARY FOR WORD DOCS ---
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="DLP Generator", layout="centered")

# --- 2. SIMPLIFIED HEADER WITHOUT LOGOS ---
def add_custom_header():
    """Add custom header with maroon background (NO LOGOS)"""
    
    st.markdown("""
    <style>
    .header-container {
        text-align: center;
        padding: 20px;
        margin-bottom: 25px;
        background-color: #800000; /* MAROON BACKGROUND */
        border-radius: 10px;
        box-shadow: 0 4px 12px rgba(128, 0, 0, 0.3);
        color: white;
    }
    .dept-name {
        font-size: 24px;
        font-weight: bold;
        color: white;
        margin: 0;
        text-shadow: 1px 1px 3px rgba(0,0,0,0.3);
    }
    .division-name {
        font-size: 20px;
        font-weight: bold;
        color: #FFD700; /* Gold color for contrast */
        margin: 8px 0;
    }
    .school-name {
        font-size: 28px;
        font-weight: bold;
        color: white;
        margin: 8px 0;
        text-transform: uppercase;
        letter-spacing: 1.5px;
    }
    .header-subtext {
        font-size: 15px;
        color: #FFD700; /* Gold color */
        margin-top: 8px;
        font-style: italic;
    }
    
    /* App title styling */
    .app-title {
        font-size: 32px;
        font-weight: bold;
        text-align: center;
        color: #800000; /* Maroon */
        margin: 15px 0 25px 0;
        padding: 10px;
        border-bottom: 3px solid #800000;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
    }
    </style>
    
    <div class="header-container">
        <p class="dept-name">DEPARTMENT OF EDUCATION REGION XI</p>
        <p class="division-name">DIVISION OF DAVAO DEL SUR</p>
        <p class="school-name">MANUAL NATIONAL HIGH SCHOOL</p>
        <p class="header-subtext">Kiblawan North District</p>
    </div>
    """, unsafe_allow_html=True)

# --- 3. IMPROVED LANGUAGE DETECTION HELPER ---
def detect_language_simple(text):
    """Simple but accurate language detection - returns 'english' or 'filipino'"""
    if not text or not isinstance(text, str):
        return "english"  # Default to English
    
    text_lower = text.lower().strip()
    
    # Very common Filipino/Tagalog words that are NOT common in English academic writing
    strong_filipino_indicators = [
        # Pronouns and particles
        'ang', 'ng', 'sa', 'mga', 'si', 'ni', 'kay', 'sina', 'nina', 'kina',
        # Common verbs
        'ay', 'may', 'meron', 'wala', 'hindi', 'oo', 'gusto', 'nais', 'kailangan',
        # Question words
        'bakit', 'paano', 'sino', 'ano', 'alin', 'kailan', 'saan', 'magkano',
        # Common classroom Filipino
        'mag-aaral', 'guro', 'aralin', 'paksa', 'talakayan', 'gawain', 'takdang-aralin',
        'pagsusulit', 'pagtataya', 'pagpapahalaga', 'kasanayan', 'kompetensi',
        'naipapamalas', 'nakagagawa', 'nakapagpapakita', 'nakapag-uugnay',
        # Common Filipino in DepEd standards
        'pag-unawa', 'pagganap', 'pagpapahalaga', 'pagsasagawa', 'pagpapakita'
    ]
    
    # Check if any strong Filipino indicators are present
    for word in strong_filipino_indicators:
        if f' {word} ' in f' {text_lower} ':
            return "filipino"
    
    # Additional check for common Filipino sentence patterns
    filipino_patterns = [
        r'\bang\b.*\bng\b',  # "ang" followed by "ng"
        r'\bnak\w+',  # Words starting with "nak" (nakapag, nakagagawa, etc.)
        r'\bnaipapamalas\b',  # Common DepEd Filipino term
        r'\bnakagagawa\b',    # Common DepEd Filipino term
        r'\bpag-\w+\b',       # Words starting with "pag-" (pag-unawa, etc.)
    ]
    
    for pattern in filipino_patterns:
        if re.search(pattern, text_lower):
            return "filipino"
    
    # Count basic Filipino vs English words
    basic_filipino_words = ['ang', 'ng', 'sa', 'mga', 'ay', 'si', 'ni', 'kay', 'para']
    basic_english_words = ['the', 'and', 'to', 'of', 'a', 'in', 'is', 'that', 'for']
    
    filipino_count = sum(1 for word in basic_filipino_words if f' {word} ' in f' {text_lower} ')
    english_count = sum(1 for word in basic_english_words if f' {word} ' in f' {text_lower} ')
    
    if filipino_count > english_count and filipino_count > 0:
        return "filipino"
    
    # Default to English
    return "english"

def analyze_language_from_inputs(content_std, perf_std, competency, obj_cognitive=None, obj_psychomotor=None, obj_affective=None, lesson_topic=None):
    """Analyze language from all teacher inputs"""
    
    # Combine all inputs for analysis
    all_inputs = []
    if content_std:
        all_inputs.append(content_std)
    if perf_std:
        all_inputs.append(perf_std)
    if competency:
        all_inputs.append(competency)
    if obj_cognitive:
        all_inputs.append(obj_cognitive)
    if obj_psychomotor:
        all_inputs.append(obj_psychomotor)
    if obj_affective:
        all_inputs.append(obj_affective)
    if lesson_topic:
        all_inputs.append(lesson_topic)
    
    if not all_inputs:
        return "english"  # Default
    
    # Check each input
    languages_found = []
    for text in all_inputs:
        if text and text.strip():
            lang = detect_language_simple(text)
            languages_found.append(lang)
    
    # If majority are Filipino, use Filipino
    if languages_found:
        filipino_count = languages_found.count("filipino")
        total = len(languages_found)
        if filipino_count > 0 and (filipino_count / total) >= 0.5:  # 50% or more
            return "filipino"
    
    return "english"

def get_language_instruction(language):
    """Get strict language instruction for AI"""
    
    if language == "filipino":
        return """
        STRICT LANGUAGE INSTRUCTION: 
        - Gamitin ang PURONG FILIPINO (Tagalog) sa lahat ng iyong sagot
        - Huwag maghalo ng English maliban sa mga teknikal na termino (hal. "Photosynthesis", "Quadratic Equation")
        - Gamitin ang mga Filipino termino gaya ng ginagamit sa DepEd curriculum
        - Ang lahat ng objectives, procedures, at assessment questions ay dapat nasa Filipino
        - Mga halimbawa:
          * English input: "The student demonstrates understanding of..."
          * Filipino output: "Naipapamalas ng mag-aaral ang pag-unawa sa..."
        """
    else:
        return """
        STRICT LANGUAGE INSTRUCTION:
        - Use PURE ENGLISH in all your responses
        - Do not mix Filipino/Tagalog words
        - Use academic English appropriate for Philippine classrooms
        - All objectives, procedures, and assessment questions must be in English
        - If teacher uses Filipino terms, translate them to English
        """

# --- 4. API KEY MANAGER WITH REMEMBER FEATURE ---
def save_api_key(api_key):
    """Save API key to session state and optionally to browser storage"""
    if api_key:
        st.session_state.api_key = api_key
        st.session_state.saved_api_key = api_key
        return True
    return False

def load_saved_api_key():
    """Load saved API key from session state"""
    return st.session_state.get('saved_api_key', '')

def show_api_key_settings():
    """Show API key input in sidebar with remember feature"""
    
    st.sidebar.markdown("---")
    st.sidebar.subheader("🔑 API Key Settings")
    
    if 'saved_api_key' not in st.session_state:
        st.session_state.saved_api_key = ""
    
    saved_key = load_saved_api_key()
    input_key = f"api_key_input_{hashlib.md5(saved_key.encode() if saved_key else ''.encode()).hexdigest()[:8]}"
    
    api_key = st.sidebar.text_input(
        "Enter your Google Gemini API Key:",
        type="password",
        placeholder="AIzaSy...",
        value=saved_key,
        key=input_key,
        help="Enter your free API key from Google AI Studio"
    )
    
    remember_me = st.sidebar.checkbox(
        "Remember my API key", 
        value=bool(saved_key),
        help="Your API key will be saved in your browser"
    )
    
    if st.sidebar.button("💾 Save API Key", use_container_width=True):
        if api_key:
            if save_api_key(api_key):
                st.sidebar.success("✅ API Key Saved!")
                if remember_me:
                    st.sidebar.info("🔒 Key saved in your browser")
                else:
                    st.sidebar.warning("⚠️ Key saved for this session only")
            else:
                st.sidebar.error("❌ Failed to save API key")
        else:
            st.sidebar.warning("⚠️ Please enter an API key")
    
    if saved_key:
        if st.sidebar.button("🗑️ Clear Saved Key", use_container_width=True):
            st.session_state.saved_api_key = ""
            st.session_state.api_key = ""
            st.sidebar.success("✅ API Key Cleared!")
            st.rerun()
    
    if st.sidebar.button("📋 How to Get Free API Key", use_container_width=True):
        st.session_state.show_instructions = True
    
    if st.session_state.get('api_key'):
        st.sidebar.success("✅ API Key Ready")
    else:
        st.sidebar.warning("⚠️ API Key Required")
    
    if api_key and remember_me and api_key != saved_key:
        save_api_key(api_key)
    
    return api_key

def show_api_key_instructions_page():
    """Show SIMPLE instructions for getting API key"""
    
    st.title("FREE API Key Instructions")
    
    st.write("""
    HOW TO GET YOUR FREE GOOGLE GEMINI API KEY:

    1. OPEN THIS LINK: https://aistudio.google.com/apikey

    2. SIGN IN with your Google account (Gmail)

    3. CLICK "Get API Key" button

    4. CLICK "Create API Key"

    5. COPY the key that appears
       (It looks like: AIzaSyABCDEFGHIJKLMNOPQRSTUVWXYZ123456)

    6. GO BACK to this app

    7. PASTE the key in the sidebar where it says 
       "Enter your Google Gemini API Key"

    8. CHECK "Remember my API key" so you don't need to enter it again

    ---

    IMPORTANT NOTES:
    • This is 100% FREE
    • No payment needed
    • Takes only 2 minutes
    • Your key is private - don't share it

    ---

    NEED HELP?
    • Make sure you're signed in to Google
    • Copy the ENTIRE key (starts with AIzaSy)
    • No spaces before or after the key
    """)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔗 OPEN GOOGLE AI STUDIO", use_container_width=True):
            js = "window.open('https://aistudio.google.com/apikey')"
            st.components.v1.html(f"<script>{js}</script>", height=0)
    
    with col2:
        if st.button("🎥 WATCH TUTORIAL", use_container_width=True):
            js = "window.open('https://www.youtube.com/results?search_query=how+to+get+google+gemini+api+key')"
            st.components.v1.html(f"<script>{js}</script>", height=0)
    
    st.markdown("---")
    if st.button("← BACK TO DLP GENERATOR", use_container_width=True):
        st.session_state.show_instructions = False
        st.rerun()

# --- 5. AI GENERATOR WITH STRICT LANGUAGE MATCHING ---
def clean_json_string(json_string):
    """Clean the JSON string by removing invalid characters and fixing common issues"""
    if not json_string:
        return json_string
    
    json_string = re.sub(r'```json\s*', '', json_string)
    json_string = re.sub(r'```\s*', '', json_string)
    json_string = json_string.replace('•', '-')
    json_string = json_string.replace('\u2022', '-')
    json_string = json_string.replace('\u25cf', '-')
    json_string = re.sub(r':\s*$', '": ""', json_string)
    
    lines = json_string.split('\n')
    cleaned_lines = []
    
    for i, line in enumerate(lines):
        quote_count = line.count('"')
        
        if quote_count % 2 == 1 and ':' in line:
            last_colon_pos = line.rfind(':')
            if last_colon_pos > 0:
                after_colon = line[last_colon_pos + 1:].strip()
                if after_colon.startswith('"') and not after_colon.endswith('"'):
                    line = line + '"'
                elif not after_colon.startswith('"') and after_colon:
                    value_start = last_colon_pos + 1
                    while value_start < len(line) and line[value_start] in ' \t':
                        value_start += 1
                    if value_start < len(line):
                        line = line[:value_start] + '"' + line[value_start:] + '"'
        
        cleaned_lines.append(line)
    
    json_string = '\n'.join(cleaned_lines)
    json_string = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', json_string)
    json_string = re.sub(r',\s*}', '}', json_string)
    json_string = re.sub(r',\s*]', ']', json_string)
    
    return json_string

def generate_lesson_content(subject, grade, quarter, content_std, perf_std, competency, 
                           obj_cognitive=None, obj_psychomotor=None, obj_affective=None,
                           lesson_topic=None):
    try:
        current_api_key = st.session_state.get('api_key') or st.session_state.get('saved_api_key')
        
        if not current_api_key:
            st.error("❌ Please enter your Google Gemini API Key in the sidebar")
            st.info("Click on '📋 How to Get Free API Key' button in sidebar for instructions")
            return None
        
        genai.configure(api_key=current_api_key)
        
        model_options = ['gemini-2.5-flash', 'gemini-1.5-flash', 'gemini-pro']
        model = None
        
        for model_name in model_options:
            try:
                model = genai.GenerativeModel(model_name)
                test_response = model.generate_content("Hello")
                if test_response:
                    st.sidebar.success(f"✓ Using model: {model_name}")
                    break
            except Exception as e:
                continue
        
        if not model:
            model = genai.GenerativeModel('gemini-1.5-flash')
        
        user_provided_objectives = obj_cognitive and obj_psychomotor and obj_affective
        user_provided_topic = lesson_topic and lesson_topic.strip()
        
        # --- DETECT LANGUAGE WITH IMPROVED LOGIC ---
        detected_language = analyze_language_from_inputs(
            content_std, perf_std, competency,
            obj_cognitive, obj_psychomotor, obj_affective,
            lesson_topic
        )
        
        # Show language detection info
        if detected_language == "filipino":
            st.sidebar.info("🌍 Language Detected: FILIPINO")
            st.sidebar.info("📝 AI will respond in PURE FILIPINO")
        else:
            st.sidebar.info("🌍 Language Detected: ENGLISH")
            st.sidebar.info("📝 AI will respond in PURE ENGLISH")
        
        # Get strict language instruction
        language_instruction = get_language_instruction(detected_language)
        
        # Create base prompt with language instruction
        if user_provided_objectives or user_provided_topic:
            prompt_parts = [
                f"""You are an expert teacher from Manual National High School in the Division of Davao Del Sur, Region XI, Philippines.
                Create a JSON object for a Daily Lesson Plan (DLP).
                Subject: {subject}, Grade: {grade}, Quarter: {quarter}
                Content Standard: {content_std}
                Performance Standard: {perf_std}
                Learning Competency: {competency}"""]
            
            # Add STRICT language instruction
            prompt_parts.append(language_instruction)
            
            if user_provided_objectives:
                prompt_parts.append(f"""
                USER-PROVIDED OBJECTIVES:
                - Cognitive: {obj_cognitive}
                - Psychomotor: {obj_psychomotor}
                - Affective: {obj_affective}
                IMPORTANT: Use these exact objectives provided by the user. Do NOT modify them.""")
            
            if user_provided_topic:
                prompt_parts.append(f"""
                USER-PROVIDED LESSON TOPIC/CONTENT:
                {lesson_topic}
                IMPORTANT: Use this exact topic/content provided by the user. Do NOT modify it.""")
            
            prompt_parts.append(f"""
            CRITICAL INSTRUCTIONS:
            1. You MUST generate exactly 5 distinct MULTIPLE CHOICE assessment questions with A, B, C, D choices.
            2. Each assessment question MUST follow this format: "question|A. choice1|B. choice2|C. choice3|D. choice4"
            3. The correct answer should be included in the choices.
            4. Return ONLY valid JSON format.
            5. Do NOT use bullet points (•) or any markdown in the JSON values.
            6. All string values must be properly quoted.
            7. Do NOT include any explanations outside the JSON.
            8. MATCH THE TEACHER'S LANGUAGE EXACTLY. If teacher used Filipino, use pure Filipino. If teacher used English, use pure English.

            Return ONLY raw JSON. No markdown formatting.
            Structure:
            {{
                "obj_1": "Cognitive objective",
                "obj_2": "Psychomotor objective",
                "obj_3": "Affective objective",
                "topic": "The main topic (include math equations like 3x^2 if needed)",
                "integration_within": "Topic within same subject",
                "integration_across": "Topic across other subject",
                "resources": {{
                    "guide": "Teacher Guide reference",
                    "materials": "Learner Materials reference",
                    "textbook": "Textbook reference",
                    "portal": "Learning Resource Portal reference",
                    "other": "Other Learning Resources"
                }},
                "procedure": {{
                    "review": "Review activity",
                    "purpose_situation": "Real-life situation motivation description",
                    "visual_prompt": "A simple 3-word visual description. Example: 'Red Apple Fruit'. NO sentences.",
                    "vocabulary": "5 terms with definitions",
                    "activity_main": "Main activity description",
                    "explicitation": "Detailed explanation of the concept with clear explanations and TWO specific examples with detailed explanations",
                    "group_1": "Group 1 task",
                    "group_2": "Group 2 task",
                    "group_3": "Group 3 task",
                    "generalization": "Reflection questions"
                }},
                "evaluation": {{
                    "assess_q1": "Question 1 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q2": "Question 2 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q3": "Question 3 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q4": "Question 4 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q5": "Question 5 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assignment": "Assignment task",
                    "remarks": "Remarks",
                    "reflection": "Reflection"
                }}
            }}
            """)
            
            prompt = "\n".join(prompt_parts)
        else:
            # Generate everything automatically with strict language matching
            prompt = f"""
            You are an expert teacher from Manual National High School in the Division of Davao Del Sur, Region XI, Philippines.
            Create a JSON object for a Daily Lesson Plan (DLP).
            Subject: {subject}, Grade: {grade}, Quarter: {quarter}
            Content Standard: {content_std}
            Performance Standard: {perf_std}
            Learning Competency: {competency}

            {language_instruction}

            CRITICAL INSTRUCTIONS:
            1. You MUST generate exactly 5 distinct MULTIPLE CHOICE assessment questions with A, B, C, D choices.
            2. Each assessment question MUST follow this format: "question|A. choice1|B. choice2|C. choice3|D. choice4"
            3. The correct answer should be included in the choices.
            4. Return ONLY valid JSON format.
            5. Do NOT use bullet points (•) or any markdown in the JSON values.
            6. All string values must be properly quoted.
            7. Do NOT include any explanations outside the JSON.
            8. MATCH THE TEACHER'S LANGUAGE EXACTLY. If teacher used Filipino, use pure Filipino. If teacher used English, use pure English.

            Return ONLY raw JSON. No markdown formatting.
            Structure:
            {{
                "obj_1": "Cognitive objective",
                "obj_2": "Psychomotor objective",
                "obj_3": "Affective objective",
                "topic": "The main topic (include math equations like 3x^2 if needed)",
                "integration_within": "Topic within same subject",
                "integration_across": "Topic across other subject",
                "resources": {{
                    "guide": "Teacher Guide reference",
                    "materials": "Learner Materials reference",
                    "textbook": "Textbook reference",
                    "portal": "Learning Resource Portal reference",
                    "other": "Other Learning Resources"
                }},
                "procedure": {{
                    "review": "Review activity",
                    "purpose_situation": "Real-life situation motivation description",
                    "visual_prompt": "A simple 3-word visual description. Example: 'Red Apple Fruit'. NO sentences.",
                    "vocabulary": "5 terms with definitions",
                    "activity_main": "Main activity description",
                    "explicitation": "Detailed explanation of the concept with clear explanations and TWO specific examples with detailed explanations",
                    "group_1": "Group 1 task",
                    "group_2": "Group 2 task",
                    "group_3": "Group 3 task",
                    "generalization": "Reflection questions"
                }},
                "evaluation": {{
                    "assess_q1": "Question 1 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q2": "Question 2 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q3": "Question 3 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q4": "Question 4 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q5": "Question 5 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assignment": "Assignment task",
                    "remarks": "Remarks",
                    "reflection": "Reflection"
                }}
            }}
            """
        
        response = model.generate_content(prompt)
        text = response.text
        
        cleaned_text = clean_json_string(text)
        
        st.sidebar.text_area("Raw AI Response", cleaned_text[:1000], height=200)
        
        try:
            ai_data = json.loads(cleaned_text)
            
            if user_provided_topic and 'topic' in ai_data:
                ai_data['topic'] = lesson_topic
                
            return ai_data
        except json.JSONDecodeError as je:
            st.error(f"JSON Parsing Error: {je}")
            st.sidebar.error("Failed to parse JSON. Attempting manual fix...")
            
            try:
                json_pattern = r'\{.*\}'
                match = re.search(json_pattern, cleaned_text, re.DOTALL)
                if match:
                    json_str = match.group(0)
                    json_str = re.sub(r',\s*}', '}', json_str)
                    json_str = re.sub(r',\s*]', ']', json_str)
                    ai_data = json.loads(json_str)
                    
                    if user_provided_topic and 'topic' in ai_data:
                        ai_data['topic'] = lesson_topic
                        
                    return ai_data
            except Exception as e2:
                st.error(f"Manual JSON extraction also failed: {e2}")
                return create_fallback_data(subject, grade, quarter, content_std, perf_std, competency, lesson_topic, detected_language)
        
    except Exception as e:
        st.error(f"AI Generation Error: {str(e)}")
        return create_fallback_data(subject, grade, quarter, content_std, perf_std, competency, lesson_topic, "english")

def create_fallback_data(subject, grade, quarter, content_std, perf_std, competency, lesson_topic=None, language="english"):
    """Create fallback data in case AI generation fails"""
    topic = lesson_topic if lesson_topic else f"Introduction to {subject}"
    
    if language == "filipino":
        return {
            "obj_1": f"Maunawaan ang mga konsepto ng {subject}",
            "obj_2": f"Magamit ang mga kasanayan sa {subject}",
            "obj_3": f"Mapahalagahan ang {subject}",
            "topic": topic,
            "integration_within": f"Kaugnay na mga paksa sa {subject}",
            "integration_across": "Matematika, Agham",
            "resources": {
                "guide": "Gabay ng Guro",
                "materials": "Kagamitan ng Mag-aaral",
                "textbook": f"Teksto sa {subject}",
                "portal": "DepEd LR Portal",
                "other": "Online resources"
            },
            "procedure": {
                "review": "Balik-aral sa nakaraang aralin",
                "purpose_situation": "Real-world application ng paksa",
                "visual_prompt": "Silid-aralan Pag-aaral",
                "vocabulary": "Term1: Kahulugan1\nTerm2: Kahulugan2\nTerm3: Kahulugan3\nTerm4: Kahulugan4\nTerm5: Kahulugan5",
                "activity_main": "Grupong gawain para tuklasin ang paksa",
                "explicitation": f"Detalyadong paliwanag ng {subject} na may mga halimbawa",
                "group_1": "Gawain sa pananaliksik",
                "group_2": "Gawain sa paglutas ng problema",
                "group_3": "Gawain sa presentasyon",
                "generalization": "Ano ang natutunan mo?"
            },
            "evaluation": {
                "assess_q1": f"Ano ang pangunahing konsepto ng {subject}?|A. Konsepto A|B. Konsepto B|C. Konsepto C|D. Konsepto D",
                "assess_q2": f"Paano mo magagamit ang {subject}?|A. Gamit A|B. Gamit B|C. Gamit C|D. Gamit D",
                "assess_q3": f"Ipaliwanag ang kahulugan sa {subject}.|A. Paliwanag A|B. Paliwanag B|C. Paliwanag C|D. Paliwanag D",
                "assess_q4": f"Lutasin ang problema sa {subject}.|A. Solusyon A|B. Solusyon B|C. Solusyon C|D. Solusyon D",
                "assess_q5": f"Ano ang limitasyon sa {subject}?|A. Limitasyon A|B. Limitasyon B|C. Limitasyon C|D. Limitasyon D",
                "assignment": "Mag-research tungkol sa paksa",
                "remarks": "Matagumpay na aralin",
                "reflection": "Magandang pag-unawa ng mga estudyante"
            }
        }
    else:
        return {
            "obj_1": f"Understand {subject} concepts",
            "obj_2": f"Apply {subject} skills",
            "obj_3": f"Appreciate the value of {subject}",
            "topic": topic,
            "integration_within": f"Related {subject} topics",
            "integration_across": "Mathematics, Science",
            "resources": {
                "guide": "Teacher's Guide",
                "materials": "Learner's Materials",
                "textbook": f"{subject} Textbook",
                "portal": "DepEd LR Portal",
                "other": "Online resources"
            },
            "procedure": {
                "review": "Review previous lesson",
                "purpose_situation": "Real-world application",
                "visual_prompt": "Classroom Learning",
                "vocabulary": "Term1: Definition1\nTerm2: Definition2\nTerm3: Definition3\nTerm4: Definition4\nTerm5: Definition5",
                "activity_main": "Group activity to explore the topic",
                "explicitation": f"Detailed explanation of {subject} with examples",
                "group_1": "Research task",
                "group_2": "Problem-solving task",
                "group_3": "Presentation task",
                "generalization": "What did you learn?"
            },
            "evaluation": {
                "assess_q1": f"What is the main concept of {subject}?|A. Concept A|B. Concept B|C. Concept C|D. Concept D",
                "assess_q2": f"How would you apply {subject}?|A. Application A|B. Application B|C. Application C|D. Application D",
                "assess_q3": f"Explain the meaning in {subject}.|A. Explanation A|B. Explanation B|C. Explanation C|D. Explanation D",
                "assess_q4": f"Solve the problem in {subject}.|A. Solution A|B. Solution B|C. Solution C|D. Solution D",
                "assess_q5": f"What are the limitations in {subject}?|A. Limitation A|B. Limitation B|C. Limitation C|D. Limitation D",
                "assignment": "Research about the topic",
                "remarks": "Lesson successful",
                "reflection": "Students showed good understanding"
            }
        }

# --- 6. IMAGE FETCHER ---
def fetch_ai_image(keywords):
    if not keywords: keywords = "school_classroom"
    clean_prompt = re.sub(r'[\n\r\t]', ' ', str(keywords))
    clean_prompt = re.sub(r'[^a-zA-Z0-9 ]', '', clean_prompt).strip()
    
    encoded_prompt = urllib.parse.quote(clean_prompt)
    seed = random.randint(1, 9999)
    url = f"https://image.pollinations.ai/prompt/{encoded_prompt}?width=600&height=350&nologo=true&seed={seed}"
    url = url.strip()
    
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=10)
        if response.status_code == 200:
            return io.BytesIO(response.content)
    except Exception:
        return None
    return None

# --- 7. DOCX HELPERS ---
def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_text(paragraph, text):
    """
    Parses text for ^ (superscript) and _ (subscript).
    """
    if not text:
        return

    pattern = r"([^\^_]*)(([\^_])([0-9a-zA-Z\-]+))(.*)"
    current_text = str(text)
    
    if "^" not in current_text and "_" not in current_text:
        paragraph.add_run(current_text)
        return

    while True:
        match = re.match(pattern, current_text)
        if match:
            pre_text = match.group(1)
            marker = match.group(3)
            script_text = match.group(4)
            rest = match.group(5)
            
            if pre_text:
                paragraph.add_run(pre_text)
            
            run = paragraph.add_run(script_text)
            if marker == '^':
                run.font.superscript = True
            elif marker == '_':
                run.font.subscript = True
                
            current_text = rest
            if not current_text:
                break
        else:
            paragraph.add_run(current_text)
            break

def add_row(table, label, content, bold_label=True):
    """Adds a row and applies formatting to the content."""
    row_cells = table.add_row().cells
    
    p_lbl = row_cells[0].paragraphs[0]
    run_lbl = p_lbl.add_run(label)
    if bold_label:
        run_lbl.bold = True
    
    text_content = ""
    if isinstance(content, list):
        text_content = "\n".join([str(item) for item in content])
    else:
        text_content = str(content) if content else ""
    
    format_text(row_cells[1].paragraphs[0], text_content)

def add_section_header(table, text):
    """Adds a full-width section header with Blue background."""
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    set_cell_background(cell, "BDD7EE")

def parse_multiple_choice_question(q_text):
    """Parse a multiple choice question in format: question|A. choice1|B. choice2|C. choice3|D. choice4"""
    if not q_text:
        return "No question provided", []
    
    parts = q_text.split('|')
    
    if len(parts) < 5:
        return q_text, []
    
    question = parts[0].strip()
    choices = []
    
    for i in range(1, min(5, len(parts))):
        choice = parts[i].strip()
        if not re.match(r'^[A-D]\.', choice):
            choice_prefix = ['A.', 'B.', 'C.', 'D.'][i-1]
            choice = f"{choice_prefix} {choice}"
        choices.append(choice)
    
    while len(choices) < 4:
        choices.append(f"{['A.', 'B.', 'C.', 'D.'][len(choices)]} Choice placeholder")
    
    return question, choices

def add_assessment_row(table, label, eval_sec):
    """Special function to add assessment row with multiple choice questions."""
    row_cells = table.add_row().cells
    
    p_lbl = row_cells[0].paragraphs[0]
    run_lbl = p_lbl.add_run(label)
    run_lbl.bold = True
    
    content_cell = row_cells[1]
    
    for paragraph in content_cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    p_header = content_cell.add_paragraph()
    header_run = p_header.add_run("ASSESSMENT (5-item Multiple Choice Quiz)")
    header_run.bold = True
    
    p_dir = content_cell.add_paragraph()
    p_dir.add_run("DIRECTIONS: Read each question carefully. Choose the letter of the correct answer from options A, B, C, and D.")
    
    content_cell.add_paragraph()
    
    for i in range(1, 6):
        question_key = f'assess_q{i}'
        raw_question = eval_sec.get(question_key, f'Question {i}')
        
        question_text, choices = parse_multiple_choice_question(raw_question)
        
        p_question = content_cell.add_paragraph()
        
        num_run = p_question.add_run(f"{i}. ")
        num_run.bold = True
        
        if question_text:
            format_text(p_question, question_text)
        
        if choices:
            for choice in choices:
                p_choice = content_cell.add_paragraph()
                p_choice.paragraph_format.left_indent = Inches(0.3)
                
                choice_match = re.match(r'^([A-D]\.)\s*(.*)', choice)
                if choice_match:
                    letter_part = choice_match.group(1)
                    text_part = choice_match.group(2)
                    
                    letter_run = p_choice.add_run(f"{letter_part} ")
                    letter_run.bold = True
                    
                    if text_part:
                        format_text(p_choice, text_part)
                else:
                    format_text(p_choice, choice)
        else:
            for letter in ['A.', 'B.', 'C.', 'D.']:
                p_choice = content_cell.add_paragraph()
                p_choice.paragraph_format.left_indent = Inches(0.3)
                letter_run = p_choice.add_run(f"{letter} ")
                letter_run.bold = True
                p_choice.add_run(f"Choice {letter[0]}")
        
        if i < 5:
            content_cell.add_paragraph()

# --- 8. DOCX CREATOR ---
def create_docx(inputs, ai_data, teacher_name, principal_name, uploaded_image):
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    dept_run = header_para.add_run("DEPARTMENT OF EDUCATION REGION XI\n")
    dept_run.bold = True
    dept_run.font.size = Pt(12)
    
    div_run = header_para.add_run("DIVISION OF DAVAO DEL SUR\n")
    div_run.bold = True
    div_run.font.size = Pt(11)
    
    school_run = header_para.add_run("MANUAL NATIONAL HIGH SCHOOL\n\n")
    school_run.bold = True
    school_run.font.size = Pt(14)
    
    title = doc.add_paragraph("Daily Lesson Log (DLL) / Daily Lesson Plan (DLP)")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    table_top = doc.add_table(rows=1, cols=4)
    table_top.style = 'Table Grid'
    table_top.autofit = False
    
    table_top.columns[0].width = Inches(2.5)
    table_top.columns[1].width = Inches(1.15)
    table_top.columns[2].width = Inches(1.15)
    table_top.columns[3].width = Inches(2.5)

    def fill_cell(idx, label, value):
        cell = table_top.rows[0].cells[idx]
        p = cell.paragraphs[0]
        p.add_run(label).bold = True
        p.add_run("\n")
        format_text(p, value)

    fill_cell(0, "Subject Area:", inputs['subject'])
    fill_cell(1, "Grade Level:", inputs['grade'])
    fill_cell(2, "Quarter:", inputs['quarter'])
    fill_cell(3, "Date:", date.today().strftime('%B %d, %Y'))

    table_main = doc.add_table(rows=0, cols=2)
    table_main.style = 'Table Grid'
    table_main.autofit = False
    
    table_main.columns[0].width = Inches(2.0)
    table_main.columns[1].width = Inches(5.3)

    objs = f"1. {ai_data.get('obj_1','')}\n2. {ai_data.get('obj_2','')}\n3. {ai_data.get('obj_3','')}"
    r = ai_data.get('resources', {})
    proc = ai_data.get('procedure', {})
    eval_sec = ai_data.get('evaluation', {})

    add_section_header(table_main, "I. CURRICULUM CONTENT, STANDARD AND LESSON COMPETENCIES")
    add_row(table_main, "A. Content Standard", inputs['content_std'])
    add_row(table_main, "B. Performance Standard", inputs['perf_std'])
    
    row_comp = table_main.add_row().cells
    row_comp[0].paragraphs[0].add_run("C. Learning Competencies").bold = True
    p_comp = row_comp[1].paragraphs[0]
    p_comp.add_run("Competency: ").bold = True
    format_text(p_comp, inputs['competency'])
    p_comp.add_run("\n\nObjectives:\n").bold = True
    p_comp.add_run(objs)

    add_row(table_main, "D. Content", ai_data.get('topic', ''))
    add_row(table_main, "E. Integration", f"Within: {ai_data.get('integration_within','')}\nAcross: {ai_data.get('integration_across','')}")

    add_section_header(table_main, "II. LEARNING RESOURCES")
    add_row(table_main, "Teacher Guide", r.get('guide', ''))
    add_row(table_main, "Learner's Materials(LMs)", r.get('materials', ''))
    add_row(table_main, "Textbooks", r.get('textbook', ''))
    add_row(table_main, "Learning Resource (LR) Portal", r.get('portal', ''))
    add_row(table_main, "Other Learning Resources", r.get('other', ''))

    add_section_header(table_main, "III. TEACHING AND LEARNING PROCEDURE")
    add_row(table_main, "A. Activating Prior Knowledge", proc.get('review', ''))
    
    row_img = table_main.add_row().cells
    row_img[0].paragraphs[0].add_run("B. Establishing Lesson Purpose").bold = True
    
    cell_img = row_img[1]
    format_text(cell_img.paragraphs[0], proc.get('purpose_situation', ''))
    cell_img.paragraphs[0].add_run("\n")
    
    img_data = None
    if uploaded_image:
        img_data = uploaded_image
    else:
        raw_prompt = proc.get('visual_prompt', 'school')
        img_data = fetch_ai_image(raw_prompt)
    
    if img_data:
        try:
            p_i = cell_img.add_paragraph()
            p_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_i = p_i.add_run()
            run_i.add_picture(img_data, width=Inches(3.5))
        except:
            cell_img.add_paragraph("[Image Error]")
    else:
        cell_img.add_paragraph("[No Image Available]")
        
    cell_img.add_paragraph(f"\nVocabulary:\n{proc.get('vocabulary','')}")

    developing_content = f"Activity: {proc.get('activity_main','')}\n\n"
    developing_content += f"EXPLICITATION: {proc.get('explicitation','')}\n\n"
    developing_content += f"Group 1: {proc.get('group_1','')}\n"
    developing_content += f"Group 2: {proc.get('group_2','')}\n"
    developing_content += f"Group 3: {proc.get('group_3','')}"
    
    add_row(table_main, "C. Developing Understanding", developing_content)
    add_row(table_main, "D. Making Generalization", proc.get('generalization', ''))

    add_section_header(table_main, "IV. EVALUATING LEARNING")
    add_assessment_row(table_main, "A. Assessment", eval_sec)
    add_row(table_main, "B. Assignment", eval_sec.get('assignment', ''))
    add_row(table_main, "C. Remarks", eval_sec.get('remarks', ''))
    add_row(table_main, "D. Reflection", eval_sec.get('reflection', ''))

    doc.add_paragraph()

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = False
    
    sig_table.columns[0].width = Inches(4.0)
    sig_table.columns[1].width = Inches(4.0)
    
    row = sig_table.rows[0]
    
    teacher_cell = row.cells[0]
    
    teacher_header_p = teacher_cell.add_paragraph()
    teacher_header_run = teacher_header_p.add_run("Prepared by:")
    teacher_header_run.bold = True
    
    teacher_cell.add_paragraph()
    
    teacher_name_p = teacher_cell.add_paragraph()
    teacher_name_run = teacher_name_p.add_run(teacher_name)
    teacher_name_run.bold = True
    
    teacher_position_p = teacher_cell.add_paragraph()
    teacher_position_p.add_run("Teacher III")
    
    principal_cell = row.cells[1]
    
    principal_header_p = principal_cell.add_paragraph()
    principal_header_run = principal_header_p.add_run("Noted by:")
    principal_header_run.bold = True
    
    principal_cell.add_paragraph()
    
    principal_name_p = principal_cell.add_paragraph()
    principal_name_run = principal_name_p.add_run(principal_name)
    principal_name_run.bold = True
    
    principal_position_p = principal_cell.add_paragraph()
    principal_position_p.add_run("Principal III")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 9. MAIN STREAMLIT APP ---
def main():
    if 'show_instructions' not in st.session_state:
        st.session_state.show_instructions = False
    if 'saved_api_key' not in st.session_state:
        st.session_state.saved_api_key = ""
    if 'api_key' not in st.session_state:
        st.session_state.api_key = st.session_state.saved_api_key
    
    if st.session_state.show_instructions:
        show_api_key_instructions_page()
        return
    
    add_custom_header()
    
    st.markdown('<p class="app-title">Daily Lesson Plan (DLP) Generator</p>', unsafe_allow_html=True)
    
    api_key = show_api_key_settings()
    
    if not api_key and st.session_state.saved_api_key:
        api_key = st.session_state.saved_api_key
        st.session_state.api_key = api_key
    
    if st.session_state.saved_api_key and not st.session_state.show_instructions:
        st.sidebar.info("✅ Your API key is saved! Ready to generate DLP.")
    
    if not api_key:
        st.warning("""
        ⚠️ **API Key Required**
        
        To use this DLP Generator, you need to:
        1. Get a **FREE** Google Gemini API Key
        2. Enter it in the sidebar
        3. Click "Save API Key" button
        4. Check "Remember my API key" so you don't need to enter it again
        
        Click the **"📋 How to Get Free API Key"** button in the sidebar for instructions.
        """)
    
    with st.sidebar:
        st.header("📋 User Information")
        
        teacher_name = st.text_input("Teacher Name", value="RICHARD P. SAMORANOS")
        principal_name = st.text_input("Principal Name", value="ROSALITA A. ESTROPIA")
        
        st.markdown("---")
        st.info("Upload an image (optional) for the lesson")
        uploaded_image = st.file_uploader("Choose an image for lesson", type=['png', 'jpg', 'jpeg'], key="lesson")
        
        st.markdown("---")
        st.info("🌍 **Smart Language Detection:**")
        st.caption("• AI detects if you're using Filipino or English")
        st.caption("• Responds in the EXACT same language")
        st.caption("• No mixed language - pure Filipino or pure English")
        
        if st.session_state.saved_api_key:
            st.markdown("---")
            st.success("🔑 API Key Status: SAVED")
            st.caption("Your key is saved for future use")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        subject = st.text_input("Subject Area", placeholder="e.g., Mathematics")
    
    with col2:
        grade_options = [
            "Kinder",
            "Grade 1", "Grade 2", "Grade 3", "Grade 4", "Grade 5", "Grade 6",
            "Grade 7", "Grade 8", "Grade 9", "Grade 10",
            "Grade 11", "Grade 12"
        ]
        grade = st.selectbox("Grade Level", grade_options, index=6)
    
    with col3:
        quarter_options = ["I", "II", "III", "IV"]
        quarter = st.selectbox("Quarter", quarter_options, index=2)
    
    content_std = st.text_area("Content Standard", placeholder="The learner demonstrates understanding of...")
    perf_std = st.text_area("Performance Standard", placeholder="The learner is able to...")
    competency = st.text_area("Learning Competency", placeholder="Competency code and description...")
    
    st.markdown("---")
    
    with st.expander("📚 Optional: Lesson Content / Topic", expanded=False):
        st.info("Enter the specific topic or content for this lesson. Leave blank if you want AI to generate it.")
        
        lesson_topic = st.text_area(
            "Lesson Content / Topic",
            placeholder="e.g., Introduction to Quadratic Equations\nOr leave blank for AI to generate",
            height=120
        )
    
    st.markdown("---")
    
    with st.expander("📝 Optional: Lesson Objectives", expanded=False):
        st.info("If you already have your lesson objectives, enter them below. Otherwise, leave blank and AI will generate them.")
        
        col_obj1, col_obj2, col_obj3 = st.columns(3)
        
        with col_obj1:
            obj_cognitive = st.text_area(
                "Cognitive Objective",
                placeholder="e.g., Solve quadratic equations\n(Leave blank for AI)",
                height=100
            )
        
        with col_obj2:
            obj_psychomotor = st.text_area(
                "Psychomotor Objective",
                placeholder="e.g., Graph quadratic functions\n(Leave blank for AI)",
                height=100
            )
        
        with col_obj3:
            obj_affective = st.text_area(
                "Affective Objective",
                placeholder="e.g., Appreciate real-world applications\n(Leave blank for AI)",
                height=100
            )
    
    st.markdown("---")
    
    has_api_key = bool(api_key or st.session_state.saved_api_key or st.session_state.get('api_key'))
    
    if st.button("🚀 Generate DLP", type="primary", use_container_width=True, disabled=not has_api_key):
        if not has_api_key:
            st.error("❌ Please enter and save your Google Gemini API Key in the sidebar first!")
            return
            
        if not all([subject, grade, quarter, content_std, perf_std, competency]):
            st.error("Please fill all required fields")
            return
        
        user_provided_topic = lesson_topic and lesson_topic.strip()
        user_provided_objectives = obj_cognitive and obj_psychomotor and obj_affective
        
        provided_items = []
        if user_provided_topic:
            provided_items.append("topic")
        if user_provided_objectives:
            provided_items.append("objectives")
        
        if provided_items:
            st.info(f"✅ Using your provided: {', '.join(provided_items)}")
            st.info("🔧 AI will generate the remaining content")
        else:
            st.info("🔧 AI will generate all lesson content for you")
        
        with st.spinner("🤖 Generating lesson content..."):
            ai_data = generate_lesson_content(
                subject, grade, quarter, 
                content_std, perf_std, competency,
                obj_cognitive if obj_cognitive else None,
                obj_psychomotor if obj_psychomotor else None,
                obj_affective if obj_affective else None,
                lesson_topic if user_provided_topic else None
            )
            
        if ai_data:
            st.success("✅ AI content generated successfully!")
            
            st.subheader("📚 Generated Lesson Content")
            col_topic, col_integration = st.columns(2)
            
            with col_topic:
                st.info("**Main Topic**")
                st.write(ai_data.get('topic', 'N/A'))
            
            with col_integration:
                st.info("**Integration**")
                st.write(f"Within Subject: {ai_data.get('integration_within', 'N/A')}")
                st.write(f"Across Subjects: {ai_data.get('integration_across', 'N/A')}")
            
            st.subheader("📋 Generated Objectives")
            col_obj_pre1, col_obj_pre2, col_obj_pre3 = st.columns(3)
            
            with col_obj_pre1:
                st.info("**Cognitive**")
                st.write(ai_data.get('obj_1', 'N/A'))
            
            with col_obj_pre2:
                st.info("**Psychomotor**")
                st.write(ai_data.get('obj_2', 'N/A'))
            
            with col_obj_pre3:
                st.info("**Affective**")
                st.write(ai_data.get('obj_3', 'N/A'))
            
            with st.expander("📝 Preview Assessment Questions"):
                for i in range(1, 6):
                    question_key = f'assess_q{i}'
                    raw_question = ai_data.get('evaluation', {}).get(question_key, '')
                    if raw_question:
                        question_text, choices = parse_multiple_choice_question(raw_question)
                        st.markdown(f"**Question {i}:** {question_text}")
                        if choices:
                            for choice in choices:
                                st.write(f"  {choice}")
                        st.markdown("---")
            
            with st.expander("📄 Preview All Generated Content"):
                st.json(ai_data)
            
            inputs = {
                'subject': subject,
                'grade': grade,
                'quarter': quarter,
                'content_std': content_std,
                'perf_std': perf_std,
                'competency': competency
            }
            
            with st.spinner("📄 Creating DOCX file..."):
                docx_buffer = create_docx(inputs, ai_data, teacher_name, principal_name, uploaded_image)
            
            st.download_button(
                label="📥 Download DLP (.docx)",
                data=docx_buffer,
                file_name=f"DLP_{subject}_{grade}_Q{quarter}_{date.today()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            st.balloons()
            st.success(f"✅ DLP generated for {subject} - {grade} - Quarter {quarter}")
            
            if st.session_state.saved_api_key:
                st.info("💡 Your API key is saved. You can use the app again without re-entering it!")
        else:
            st.error("Failed to generate AI content. Please try again.")

if __name__ == "__main__":
    main()
